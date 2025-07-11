from fastapi import APIRouter, HTTPException, status
from fastapi.responses import FileResponse
from base_requests import GenerateContentRequest, GenerateContentResponse, ScanJobRequest, ScanJobResponse
from util.system_prompt import agg_prompt
from util.llm_factory import LLMFactory
from jd_scanning import invoke_llm_with_reflection
import json
import re
import logging
from test_run import generate_job_description
from util.utility import Utility
from docx import Document
import os
import uuid

def scan_job_description(jd: str, config_path: str, count: int = 3) -> dict:
    # Use the provided config_path, fallback to default if None
    try:
        with open(config_path, 'r') as f:
            format_schema = f.read()
    except FileNotFoundError:
        raise ValueError(f"Configuration file not found: {config_path}")
    except Exception as e:
        raise ValueError(f"Error reading configuration file: {e}")
    reflections = invoke_llm_with_reflection(jd, format_schema)

    aggregation_prompt = agg_prompt.format(reflected_outputs=json.dumps(reflections, indent=2)).strip()
    final_result = LLMFactory.create_llm_instance(temperature=0.2, local_llm=False).invoke(aggregation_prompt)

    # Try to extract JSON using regex if LLM adds extra text
    match = re.search(r"\{.*\}", final_result.content, re.DOTALL)
    if match:
        cleaned_json = match.group(0)
        return json.loads(cleaned_json)
    else:
        raise ValueError("LLM response did not contain valid JSON.")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

api_router = APIRouter(tags=["HR API Services"])


@api_router.post(
    "/generate",
    response_model=GenerateContentResponse,
    status_code=status.HTTP_200_OK,
    summary="Generate JD and return download link",
    responses={
        400: {"description": "Invalid Request"},
        422: {"description": "Unprocessable Entity"},
        500: {"description": "Internal Server Error"},
    },
)
async def generate_content(request: GenerateContentRequest):
    try:
        logger.info("Generating content with provided job parameters...")
        input_data = request.model_dump()
        job_description = generate_job_description(**input_data)
        job_description = Utility.clean_text(job_description, preserve_paragraphs=True)

        if not job_description:
            raise HTTPException(status_code=500, detail="Failed to generate job description")

        # Save as .docx
        doc = Document()
        doc.add_heading('Job Description', 0)
        doc.add_paragraph(job_description)

        output_dir = "generated_docs"
        os.makedirs(output_dir, exist_ok=True)
        file_id = str(uuid.uuid4())
        output_path = os.path.join(output_dir, f"job_description_{file_id}.docx")
        doc.save(output_path)

        logger.info(f"Job description saved to {output_path}")

        download_url = f"/api/v1/download/{file_id}"  # Assuming prefix is /api/v1

        return {
            "status": "success",
            "message": "Content generated successfully",
            "data": {
                "job_description": job_description,
                "download_url": download_url
            }
        }
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error generating content: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error generating content: {str(e)}"
        )

@api_router.post(
    "/scan",
    response_model=ScanJobResponse,
    status_code=status.HTTP_200_OK,
    summary="Scan JD and return structured analysis",
    responses={
        400: {"description": "Invalid Job Description"},
        422: {"description": "Unprocessable Job Description"},
        500: {"description": "Internal Server Error"},
    },
)
async def scan_jd(request: ScanJobRequest):
    try:
        logger.info("Scanning job description for structured analysis...")
        summary = scan_job_description(
            jd=request.job_description,
            config_path=request.config_file_path,
            count=getattr(request, 'llm_reflection_count', 3)
        )
        return ScanJobResponse(
            status="success", message="Content generated successfully", data=summary
        )
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Error generating content: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error generating content: {str(e)}"
        )


@api_router.get(
    "/download/{file_id}",
    summary="Download the generated DOCX file",
    response_class=FileResponse
)
async def download_docx(file_id: str):
    output_dir = "generated_docs"
    file_path = os.path.join(output_dir, f"job_description_{file_id}.docx")
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="job_description.docx"
    )
